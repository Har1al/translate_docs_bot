from docx import Document
import asyncio
from openai import OpenAI
from environs import Env
from services.call_model import call_model_for_translations

env = Env()
env.read_env()



async def _spinner_edit(status_message, base_text="Отправляю запрос модели"):
    dots = 0
    try:
        while True:
            dots = (dots % 3) + 1
            try:
                await status_message.edit_text(f"{base_text}{'.' * dots}")
            except Exception:
                pass
            await asyncio.sleep(0.7)
    except asyncio.CancelledError:
        raise



def is_text_run(run):
    """Run содержит текст, а не картинку"""
    try:
        if run._element.xpath('.//w:drawing'):
            return False
    except Exception:
        pass
    return bool(run.text and run.text.strip())


def collect_texts(doc: Document):
    texts = set()

    def collect_from_paragraphs(paragraphs):
        for p in paragraphs:
            for r in p.runs:
                if is_text_run(r):
                    texts.add(r.text.strip())

    collect_from_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                collect_from_paragraphs(cell.paragraphs)

    for section in doc.sections:
        collect_from_paragraphs(section.header.paragraphs)
        collect_from_paragraphs(section.footer.paragraphs)

    return list(texts)



def replace_runs(paragraph, mapping):
    for run in paragraph.runs:
        if not is_text_run(run):
            continue

        original = run.text.strip()
        if original in mapping:
            run.text = mapping[original]


def process_document(doc: Document, mapping: dict):

    for p in doc.paragraphs:
        replace_runs(p, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_runs(p, mapping)

    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_runs(p, mapping)
        for p in section.footer.paragraphs:
            replace_runs(p, mapping)



async def from_eng_to_rus_docx(path, message, save_path):

    doc = Document(path)
    texts = collect_texts(doc)

    if not texts:
        await message.answer("В документе нет текста для перевода.")
        return

    await message.answer(f"Найдено {len(texts)} текстовых фрагментов для перевода.")

    client = OpenAI(api_key=env("API_KEY"))

    status_msg = await message.answer("Отправляю запрос модели")
    spinner_task = asyncio.create_task(_spinner_edit(status_msg))

    try:
        mapping = await asyncio.to_thread(
            call_model_for_translations,
            client,
            texts
        )

        spinner_task.cancel()
        try:
            await spinner_task
        except asyncio.CancelledError:
            pass

        if not isinstance(mapping, dict):
            await status_msg.edit_text("❌ Модель вернула некорректный ответ.")
            return

        await status_msg.edit_text(f"✅ Получено {len(mapping)} переводов.")

        process_document(doc, mapping)
        doc.save(save_path)

        await message.answer("📄 Документ успешно переведён.")

    except Exception as e:
        spinner_task.cancel()
        try:
            await spinner_task
        except asyncio.CancelledError:
            pass
        await status_msg.edit_text("❌ Ошибка при обработке документа.")
        print("[ERROR]", e)
